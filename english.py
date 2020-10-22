from docx import *
from docx.shared import Pt
from docx.shared import Inches
#Generate a template for MLA format essay.
def generateDoc():
    document = Document()

    #Line spacing.
    paragraph = document.add_paragraph()
    paragraph_run = paragraph.add_run()
    format = paragraph.paragraph_format
    format.line_spacing = 2
    #Font
    font = paragraph_run.font()
    font.name = "Times New Roman"
    font.size = Pt(12)

    #Page Size
    sections = document.sections
    section = sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)

    #Margins
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.top_margin = Inches(1)
    # Header that numbers pages in top right hand corner and your last name. Ryan 1, Ryan 2, etc.
    # If you have endnotes include it before works cited, entitle it notes and center it.

    #Your name, Instructors name, course, date. Double spaced text.
    #Double space and center the title.
    
    document.save("mlaessay.docx")

#Generate a template for APA format essay.

#Change format to MLA
#Change format to APA

#Indeed